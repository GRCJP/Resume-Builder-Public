#!/usr/bin/env python3
"""Generate optimized GRC resume in DOCX and PDF formats"""

import sys
from pathlib import Path

def install_packages():
    """Install required packages"""
    import subprocess
    packages = ['python-docx', 'reportlab']
    for package in packages:
        try:
            __import__(package.replace('-', '_'))
        except ImportError:
            print(f"Installing {package}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])

install_packages()

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER

def create_docx_resume():
    """Create DOCX version of optimized resume"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    # Header - Name
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run('JONATHAN L. PEREZ, MS, CISSP, CISM')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Contact Info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact.add_run('Aldie, VA | 703-843-6887 | jleepe@outlook.com | linkedin.com/in/cyberjp | securitybyjp.com')
    run.font.size = Pt(10)
    
    # Clearance
    clearance = doc.add_paragraph()
    clearance.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = clearance.add_run('Active Secret Clearance')
    run.font.size = Pt(10)
    run.font.italic = True
    
    doc.add_paragraph()  # Spacer
    
    # ROADMAP OF IMPACT
    heading = doc.add_paragraph()
    run = heading.add_run('ROADMAP OF IMPACT')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    impacts = [
        'Reduced enterprise vulnerabilities from 70,000 → 12,000 across 10,000+ assets through automated prioritization and remediation workflows',
        'Engineered continuous compliance pipelines using AWS Config, Lambda, EventBridge, and Security Hub to automatically collect evidence and detect drift',
        'Automated SSP→OSCAL→ARCAMPE documentation workflows, reducing manual cycles by 70% and improving audit accuracy',
        'Built vulnerability → POA&M automation engine ingesting Wiz/Qualys/Tenable outputs, cutting manual review hours by 30%',
        'Improved CIS/STIG compliance to 90%+ across 7,500+ cloud assets (MD THINK/MD Benefits)',
        'Delivered ISSO enablement programs reducing artifact submission errors by 40% through improved workflows and training'
    ]
    
    for impact in impacts:
        p = doc.add_paragraph(impact, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    doc.add_paragraph()  # Spacer
    
    # PROFESSIONAL SUMMARY
    heading = doc.add_paragraph()
    run = heading.add_run('PROFESSIONAL SUMMARY')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    summary_text = """GRC Engineer and Senior Assessor with 10+ years of experience spanning NIST RMF, FedRAMP, SOC 2, ISO 27001, PCI DSS, and HIPAA compliance across Federal and State environments. Proven expertise merging ISSO-level governance with engineering execution to automate evidence collection, strengthen vulnerability reduction, and implement cloud-native compliance workflows.

Experienced in full RMF lifecycle (SSP, SAR, POA&M, SIA), vulnerability analytics, CI/CD security validation, SIEM evidence review, and risk register automation. Track record delivering measurable improvements in audit readiness, compliance efficiency, and remediation velocity. Skilled at translating complex technical requirements into actionable risk management strategies for both technical and executive stakeholders."""
    
    p = doc.add_paragraph(summary_text)
    for run in p.runs:
        run.font.size = Pt(10)
    
    doc.add_paragraph()  # Spacer
    
    # CORE COMPETENCIES
    heading = doc.add_paragraph()
    run = heading.add_run('CORE COMPETENCIES')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    competencies = [
        ('GRC Frameworks & Compliance:', 'NIST RMF, NIST 800-53/800-171/800-30/800-37/800-18, FedRAMP, FISMA, SOC 2, ISO 27001/27002, PCI DSS 4.0, HIPAA, GDPR, IRS Pub 1075, CMS ARCAMPE, NIST CSF, CIS Controls, DISA STIGs, COBIT'),
        ('Cloud & Automation:', 'AWS Config, Lambda, EventBridge, Security Hub, Systems Manager, GuardDuty, ECR, Python, OSCAL, YAML, API integration, CI/CD security validation'),
        ('GRC Platforms & Tools:', 'ServiceNow GRC, DOJ CSAM, RegScale, Jira Automation, Confluence, API-based evidence pipelines, automated risk registers'),
        ('Vulnerability Management:', 'Wiz, Qualys, Tenable, CrowdStrike Falcon, AWS Security Hub, Burp Suite, Fortify, Checkmarx, CVSSv4, EPSS, MITRE ATT&CK, container scanning, risk-based triage'),
        ('Security Operations:', 'SIEM analysis (Splunk, Security Hub, GuardDuty), incident response validation, log correlation, continuous monitoring, threat assessment'),
        ('DevSecOps & CI/CD:', 'GitHub Actions, Semgrep, Trivy, container image scanning, baseline validation, automated security testing'),
        ('Governance & Documentation:', 'System Security Plans (SSP), POA&Ms, Security Assessment Reports (SAR), Security Authorization Packages (SAP), risk assessments, policy development, audit preparation')
    ]
    
    for category, details in competencies:
        p = doc.add_paragraph()
        run = p.add_run(category + ' ')
        run.font.size = Pt(10)
        run.font.bold = True
        run = p.add_run(details)
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_paragraph()  # Spacer
    
    # PROFESSIONAL EXPERIENCE
    heading = doc.add_paragraph()
    run = heading.add_run('PROFESSIONAL EXPERIENCE')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Job 1 - Assurit
    job_title = doc.add_paragraph()
    run = job_title.add_run('ASSURIT')
    run.font.size = Pt(11)
    run.font.bold = True
    
    job_details = doc.add_paragraph()
    run = job_details.add_run('Senior Cybersecurity Specialist / Deputy Project Manager')
    run.font.size = Pt(10)
    run.font.bold = True
    run = job_details.add_run(' | September 2018 – Present | Washington, D.C.')
    run.font.size = Pt(10)
    run.font.italic = True
    
    assurit_bullets = [
        'Direct enterprise vulnerability management program across 10,000+ assets, leading cross-functional teams in remediating vulnerabilities and partnering with engineers to resolve complex issues not easily fixed through patches or workarounds',
        'Serve as Subject Matter Expert on NIST RMF, continuous monitoring, security control assessments, and vendor risk reviews, ensuring compliance with federal cybersecurity regulations including FedRAMP, FISMA, SOC 2, ISO 27001, and PCI DSS',
        'Lead weekly vulnerability management sessions to assess threat landscape relevance, prioritize findings by risk and asset criticality, validate false positives, and coordinate retesting and remediation with stakeholders using CVSS-based prioritization',
        'Engineered automated remediation routing and prioritization logic, reducing vulnerabilities from 70,000 to 12,000 through control-mapped workflows and systematic risk-based triage',
        'Built AWS-native continuous compliance pipelines using Config, Lambda, EventBridge, and Security Hub to detect drift, collect evidence automatically, and sync findings to Jira/ServiceNow for remediation tracking',
        'Automated RMF evidence pipelines by transforming Wiz/Qualys/Tenable data into normalized, control-mapped findings, reducing manual review time by 30% and improving audit accuracy',
        'Developed SSP→OSCAL→ARCAMPE conversion workflows using Python and RegScale, reducing documentation cycles by 70% and increasing artifact consistency across multiple frameworks',
        'Conducted comprehensive security assessments for MD THINK, MD Benefits, FEC, and USPTO, covering NIST 800-53/800-171, FedRAMP, ISO 27001, SOC 2, IRS Pub 1075, and CMS ARCAMPE requirements',
        'Validated CI/CD pipeline security controls including ECR scanning, Wiz SCA, and SSM patching, aligning build processes with NIST 800-53 and FedRAMP baseline requirements',
        'Performed SIEM analysis leveraging Security Hub, GuardDuty, and Splunk to validate detection controls, triage audit-relevant events, and support incident response documentation',
        'Designed Jira-based risk register automation improving risk scoring, tracking workflows, and continuous monitoring updates across multiple compliance frameworks',
        'Delivered ISSO enablement programs that improved artifact quality and reduced submission errors by 40% through standardized workflows and targeted training',
        'Managed POA&M lifecycle from identification through closure, tracking remediation progress and ensuring timely resolution of audit findings',
        'Coordinated with Legal, IT, and Compliance teams on governance initiatives, translating technical security requirements into business-aligned risk management strategies'
    ]
    
    for bullet in assurit_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    # Selected Achievements - Assurit
    ach_heading = doc.add_paragraph()
    run = ach_heading.add_run('Selected Achievements:')
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.italic = True
    ach_heading.paragraph_format.left_indent = Inches(0.25)
    
    assurit_achievements = [
        'Improved CIS/STIG compliance to 90%+ across 7,500+ cloud assets for MD THINK and MD Benefits',
        'Reduced critical vulnerabilities by 65% in one quarter at MD DHS through targeted remediation campaigns',
        'Implemented AWS drift detection workflows that cut control validation time by weeks per assessment cycle',
        'Achieved zero material weaknesses in SOC 2 audits through proactive control testing and remediation tracking'
    ]
    
    for ach in assurit_achievements:
        p = doc.add_paragraph(ach, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
        for run in p.runs:
            run.font.size = Pt(10)
    
    doc.add_paragraph()  # Spacer
    
    # Job 2 - Various Federal Contractors
    job_title = doc.add_paragraph()
    run = job_title.add_run('USPTO / SEC / DHS (Various Federal Contractors)')
    run.font.size = Pt(11)
    run.font.bold = True
    
    job_details = doc.add_paragraph()
    run = job_details.add_run('Senior Information Security Analyst / Information System Security Officer (ISSO)')
    run.font.size = Pt(10)
    run.font.bold = True
    run = job_details.add_run(' | April 2014 – September 2018 | Washington, D.C.')
    run.font.size = Pt(10)
    run.font.italic = True
    
    isso_bullets = [
        'Ensured compliance with NIST, FedRAMP, and FISMA standards, supporting audit readiness and security control testing using NIST SP 800 series, FIPS, and ISO 27000 frameworks',
        'Performed full RMF lifecycle activities including System Security Plan (SSP) development, Security Assessment Reports (SAR), POA&M management, Security Impact Analysis (SIA), and PIA/PTA documentation',
        'Reviewed federal and COTS configuration documentation to assess potential security vulnerabilities and risks, ensuring continuous alignment with regulatory requirements',
        'Communicated priorities to stakeholders and presented bi-monthly briefings to review security issues and prioritize remediation efforts across multiple federal agencies',
        'Implemented and managed secure baseline configurations for Windows platforms in alignment with CIS Benchmarks, USGCB, and DISA STIGs, reinforcing system security and standardization',
        'Led change management initiatives, reviewing change requests and driving product adoption and end-user training while conducting security impact analysis',
        'Managed vulnerability and compliance scanning in support of independent verification and validation (IV&V), audits, and incident response across UNIX, Cisco, and Windows platforms',
        'Served as Subject Matter Expert on IDS/IPS, port and vulnerability scanners, network detection, and security requirements for high-value federal assets',
        'Coordinated cross-platform troubleshooting efforts with OS Support, Database, Networking, and VMware teams to resolve problematic scans and connectivity issues',
        'Performed code-embedded security reviews for SEC development teams, reducing release-blocking vulnerabilities through early detection and remediation',
        'Maintained CSAM risk registers, updated mitigation plans, and increased remediation tracking accuracy across multiple federal systems',
        'Supported SOC 2 and ISO 27001 evidence collection and technical control mapping for compliance validation',
        'Strengthened incident response readiness by correlating SIEM logs with endpoint and cloud telemetry for audit verification'
    ]
    
    for bullet in isso_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    # Selected Achievements - ISSO
    ach_heading = doc.add_paragraph()
    run = ach_heading.add_run('Selected Achievements:')
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.italic = True
    ach_heading.paragraph_format.left_indent = Inches(0.25)
    
    isso_achievements = [
        'Conducted comprehensive vulnerability assessments to map NIST controls, strengthen organizational compliance, and protect intellectual property for the United States Patent and Trademark Office',
        'Provided in-depth source code analysis during short-term contract with Securities and Exchange Commission, ensuring secure coding practices within strict project timelines',
        'Developed resolution policies that ensured rapid response to highly visible vulnerabilities identified during continuous monitoring, significantly improving mitigation turnaround times',
        'Created automated scanning workflow that increased productivity by 20% and enhanced scan data accuracy',
        'Designed and implemented After-Action Report (AAR) process to enhance tracking of automated scans and standardize reporting across multiple teams',
        'Tailored audit policies and scanning tools to minimize unnecessary data output, streamlining analysis and improving actionable insights for remediation',
        'Reduced false positives by 25% through refined scanning configurations and validation processes'
    ]
    
    for ach in isso_achievements:
        p = doc.add_paragraph(ach, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
        for run in p.runs:
            run.font.size = Pt(10)
    
    # Page break before next section
    doc.add_page_break()
    
    # FEATURED PROJECTS
    heading = doc.add_paragraph()
    run = heading.add_run('FEATURED PROJECTS & AUTOMATION INITIATIVES')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    projects = [
        ('AWS Config → Jira Risk Register Automation', 'Built Config/Lambda/EventBridge workflow to detect drift, parse findings, and push structured events into Jira for continuous compliance tracking and automated risk scoring'),
        ('Vulnerability Metrics Automation Engine', 'Developed Python scripts to normalize Wiz/Qualys/Tenable exports and generate MTTR, aging, and exposure dashboards for executive reporting'),
        ('SSP→OSCAL→ARCAMPE Conversion Pipeline', 'Automated documentation workflows reducing SSP update cycles by 70% and increasing artifact consistency across multiple compliance frameworks'),
        ('Control Harmonization Framework', 'Mapped NIST 800-53, SOC 2, ISO 27001, IRS Pub 1075, and CMS ARCAMPE to unify controls, evidence paths, and narratives for streamlined assessments'),
        ('POA&M Automation Engine', 'Engineered vulnerability-to-POA&M automation ingesting scanner outputs and reducing manual review hours by 30%')
    ]
    
    for project_name, description in projects:
        p = doc.add_paragraph()
        run = p.add_run(project_name + ': ')
        run.font.size = Pt(10)
        run.font.bold = True
        run = p.add_run(description)
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_paragraph()  # Spacer
    
    # EDUCATION
    heading = doc.add_paragraph()
    run = heading.add_run('EDUCATION')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    education = [
        'Master of Science (MS), Cybersecurity – University of Maryland Global Campus',
        'Bachelor of Fine Arts (BFA) – Savannah College of Art and Design'
    ]
    
    for edu in education:
        p = doc.add_paragraph(edu, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    doc.add_paragraph()  # Spacer
    
    # CERTIFICATIONS
    heading = doc.add_paragraph()
    run = heading.add_run('CERTIFICATIONS')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    certs_text = 'Certified Information Systems Security Professional (CISSP) • Certified Information Security Manager (CISM) • Certified in Governance, Risk and Compliance (CGRC) • AWS Solutions Architect – Associate • Certificate of Cloud Security Knowledge (CCSK) • Certified Ethical Hacker (CEH) • Certified AI Security Specialist (CAISS) • CompTIA Security+ (SEC+) • CompTIA Network+ (NET+) • Certificate of Competence in Zero Trust (CCZT) • Trusted AI Safety Expert (TAISE) • CISA AES Assessment Lead (AL) • CISA AES High Value Asset Technical Lead (TL)'
    
    p = doc.add_paragraph(certs_text)
    for run in p.runs:
        run.font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.25)
    
    # Save DOCX
    output_path = Path(__file__).parent / "Resume" / "Jonathan_Perez_GRC_Resume_OPTIMIZED.docx"
    doc.save(output_path)
    print(f"✅ DOCX saved: {output_path}")
    return output_path

def create_pdf_resume():
    """Create PDF version (simplified for now)"""
    output_path = Path(__file__).parent / "Resume" / "Jonathan_Perez_GRC_Resume_OPTIMIZED.pdf"
    
    pdf = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        rightMargin=0.7*inch,
        leftMargin=0.7*inch,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch
    )
    
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor='#003366',
        spaceAfter=6,
        alignment=TA_CENTER
    )
    
    contact_style = ParagraphStyle(
        'Contact',
        parent=styles['Normal'],
        fontSize=10,
        alignment=TA_CENTER,
        spaceAfter=12
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=12,
        textColor='#003366',
        spaceAfter=6,
        spaceBefore=12
    )
    
    # Title
    story.append(Paragraph('<b>JONATHAN L. PEREZ, MS, CISSP, CISM</b>', title_style))
    story.append(Paragraph('Aldie, VA | 703-843-6887 | jleepe@outlook.com | linkedin.com/in/cyberjp | securitybyjp.com', contact_style))
    story.append(Paragraph('<i>Active Secret Clearance</i>', contact_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Note about PDF
    note = Paragraph(
        '<b>Note:</b> For best formatting, please use the DOCX version. '
        'This PDF is a simplified version for quick reference.',
        styles['Normal']
    )
    story.append(note)
    story.append(Spacer(1, 0.2*inch))
    
    # Summary
    story.append(Paragraph('<b>PROFESSIONAL SUMMARY</b>', heading_style))
    summary = Paragraph(
        'GRC Engineer and Senior Assessor with 10+ years of experience spanning NIST RMF, FedRAMP, SOC 2, '
        'ISO 27001, PCI DSS, and HIPAA compliance across Federal and State environments. Proven expertise '
        'merging ISSO-level governance with engineering execution to automate evidence collection, strengthen '
        'vulnerability reduction, and implement cloud-native compliance workflows.',
        styles['Normal']
    )
    story.append(summary)
    
    # Build PDF
    pdf.build(story)
    print(f"✅ PDF saved: {output_path}")
    print("\n⚠️  Note: The PDF is a simplified version. Use the DOCX for best formatting and ATS compatibility.")
    return output_path

if __name__ == "__main__":
    print("Generating optimized GRC resume...\n")
    docx_path = create_docx_resume()
    pdf_path = create_pdf_resume()
    print(f"\n✅ Resume generation complete!")
    print(f"\nFiles created:")
    print(f"  - {docx_path.name} (RECOMMENDED - Use this for applications)")
    print(f"  - {pdf_path.name} (Simplified reference version)")
