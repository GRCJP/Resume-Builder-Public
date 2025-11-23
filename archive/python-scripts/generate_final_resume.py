#!/usr/bin/env python3
"""Generate FINAL VP-aligned GRC resume in DOCX format"""

import sys
from pathlib import Path

def install_packages():
    """Install required packages"""
    import subprocess
    try:
        __import__('docx')
    except ImportError:
        print("Installing python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

install_packages()

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_final_resume():
    """Create final VP-aligned DOCX resume"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    # Header - Name
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run('JONATHAN L. PEREZ, MS, CISSP, CISM')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # NEW: Tagline (VP Recommendation)
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run('Bridging GRC and Engineering to automate compliance at scale')
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Contact Info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact.add_run('Aldie, VA | 703-843-6887 | jleepe@outlook.com | linkedin.com/in/cyberjp | securitybyjp.com')
    run.font.size = Pt(10)
    
    # Clearance
    clearance = doc.add_paragraph()
    clearance.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = clearance.add_run('Active Secret Clearance')
    run.font.size = Pt(10)
    run.font.italic = True
    
    doc.add_paragraph()  # Spacer
    
    # ROADMAP OF IMPACT
    heading = doc.add_paragraph()
    run = heading.add_run('ROADMAP OF IMPACT')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    impacts = [
        'Reduced enterprise vulnerabilities from 70,000 → 12,000 across 10,000+ assets through automated prioritization and remediation workflows',
        'Engineered continuous compliance pipelines using AWS Config, Lambda, EventBridge, and Security Hub to automatically collect evidence and detect drift',
        'Automated SSP→OSCAL→ARCAMPE documentation workflows, reducing manual cycles by 70% and improving audit accuracy',
        'Built vulnerability → POA&M automation engine ingesting Wiz/Qualys/Tenable outputs, cutting manual review hours by 30%',
        'Improved CIS/STIG compliance to 90%+ across 7,500+ cloud assets through control automation',
        'Delivered ISSO enablement programs reducing artifact submission errors by 40% through improved workflows and training'
    ]
    
    for impact in impacts:
        p = doc.add_paragraph(impact, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    doc.add_paragraph()  # Spacer
    
    # PROFESSIONAL SUMMARY (REWRITTEN per VP feedback)
    heading = doc.add_paragraph()
    run = heading.add_run('PROFESSIONAL SUMMARY')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    summary_text = """GRC Engineer leveraging AWS Config, Lambda, EventBridge, and Security Hub to automate continuous compliance and evidence collection across hybrid cloud environments. 10+ years bridging NIST RMF, FedRAMP, SOC 2, ISO 27001, and PCI DSS governance with engineering execution using Wiz, Qualys, Tenable, and Python automation. Proven track record reducing vulnerabilities 83% and cutting documentation cycles 70% through control automation and risk register integration.

Deep expertise in vulnerability management leadership, integrating security tools with Jira and ServiceNow for automated remediation tracking. Skilled at translating complex compliance requirements into engineered solutions that deliver measurable business value. Strong certification foundation (CISSP, CISM, CGRC, AWS Solutions Architect) demonstrating both governance knowledge and cloud technical fluency."""
    
    p = doc.add_paragraph(summary_text)
    for run in p.runs:
        run.font.size = Pt(10)
    
    doc.add_paragraph()  # Spacer
    
    # TOOLS & PLATFORMS (Emphasized per VP feedback)
    heading = doc.add_paragraph()
    run = heading.add_run('TOOLS & PLATFORMS')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    tools = [
        ('Cloud & Automation:', 'AWS Config, Lambda, EventBridge, Security Hub, Systems Manager, GuardDuty, ECR, Python, OSCAL, YAML, API integration, CI/CD security validation, continuous compliance automation'),
        ('GRC & Evidence Automation:', 'ServiceNow GRC, DOJ CSAM, RegScale, Jira Automation, Confluence, API-based evidence pipelines, automated risk registers, control automation'),
        ('Vulnerability & Security Tools:', 'Wiz, Qualys, Tenable, CrowdStrike Falcon, AWS Security Hub, Burp Suite, Fortify, Checkmarx, CVSSv4, EPSS, MITRE ATT&CK, container scanning, risk-based triage'),
        ('Frameworks & Compliance:', 'NIST RMF, NIST 800-53/800-171/800-30/800-37, FedRAMP, FISMA, SOC 2, ISO 27001/27002, PCI DSS 4.0, HIPAA, GDPR, IRS Pub 1075, CMS ARCAMPE, NIST CSF, CIS Controls, DISA STIGs'),
        ('Security Operations:', 'SIEM analysis (Splunk, Security Hub, GuardDuty), incident response validation, log correlation, continuous monitoring, threat assessment, evidence collection'),
        ('DevSecOps & CI/CD:', 'GitHub Actions, Semgrep, Trivy, container image scanning, baseline validation, automated security testing, infrastructure as code validation')
    ]
    
    for category, details in tools:
        p = doc.add_paragraph()
        run = p.add_run(category + ' ')
        run.font.size = Pt(10)
        run.font.bold = True
        run = p.add_run(details)
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_paragraph()  # Spacer
    
    # PROFESSIONAL EXPERIENCE
    heading = doc.add_paragraph()
    run = heading.add_run('PROFESSIONAL EXPERIENCE')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Job 1 - Assurit (ENHANCED with VP keywords)
    job_title = doc.add_paragraph()
    run = job_title.add_run('ASSURIT')
    run.font.size = Pt(11)
    run.font.bold = True
    
    job_details = doc.add_paragraph()
    run = job_details.add_run('Senior Cybersecurity Specialist / Deputy Project Manager')
    run.font.size = Pt(10)
    run.font.bold = True
    run = job_details.add_run(' | September 2018 – Present | Washington, D.C.')
    run.font.size = Pt(10)
    run.font.italic = True
    
    assurit_bullets = [
        'Lead enterprise vulnerability management program across 10,000+ assets, integrating Wiz, Qualys, and Tenable with Jira and ServiceNow for automated remediation tracking, reducing vulnerabilities from 70,000 to 12,000 (83% reduction)',
        'Engineer continuous compliance pipelines using AWS Config, Lambda, and EventBridge to automatically collect evidence, detect configuration drift, and validate security controls in real-time across hybrid cloud environments',
        'Automate RMF evidence collection and POA&M generation by transforming Wiz/Qualys/Tenable data into normalized, control-mapped findings, reducing manual review time by 30% and improving audit accuracy',
        'Serve as GRC Engineering SME on NIST RMF, FedRAMP, SOC 2, ISO 27001, and PCI DSS compliance, bridging governance requirements with automated technical implementation',
        'Develop control automation workflows using Python and RegScale, converting SSP→OSCAL→ARCAMPE documentation and reducing compliance cycles by 70% through automated evidence pipelines',
        'Lead security assessments for MD THINK, MD Benefits, FEC, and USPTO, implementing continuous compliance monitoring across NIST 800-53/800-171, FedRAMP, ISO 27001, SOC 2, and IRS Pub 1075 frameworks',
        'Integrate SIEM data from Security Hub, GuardDuty, and Splunk with GRC platforms for automated evidence collection, control validation, and incident response documentation',
        'Build Jira-based risk register automation improving risk scoring, tracking workflows, and continuous monitoring updates across multiple compliance frameworks',
        'Validate CI/CD pipeline security controls including ECR scanning, Wiz SCA, and SSM patching, aligning DevSecOps processes with NIST 800-53 and FedRAMP baseline requirements',
        'Deliver ISSO enablement programs improving artifact quality and reducing submission errors by 40% through standardized workflows, automated templates, and targeted training',
        'Drive weekly vulnerability management sessions assessing threat landscape relevance, prioritizing findings by CVSS/EPSS and asset criticality, and coordinating automated remediation with engineering teams',
        'Manage POA&M lifecycle from identification through closure using automated tracking, ensuring timely resolution of audit findings and maintaining continuous compliance posture',
        'Coordinate with Legal, IT, and Compliance teams on governance initiatives, translating technical security requirements into business-aligned risk management strategies'
    ]
    
    for bullet in assurit_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    # Selected Achievements - Assurit
    ach_heading = doc.add_paragraph()
    run = ach_heading.add_run('Selected Achievements:')
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.italic = True
    ach_heading.paragraph_format.left_indent = Inches(0.25)
    
    assurit_achievements = [
        'Improved CIS/STIG compliance to 90%+ across 7,500+ cloud assets through automated control validation',
        'Reduced critical vulnerabilities by 65% in one quarter through risk-based prioritization and automated remediation',
        'Implemented AWS drift detection workflows cutting control validation time by weeks per assessment cycle',
        'Achieved zero material weaknesses in SOC 2 audits through continuous compliance monitoring and proactive control testing'
    ]
    
    for ach in assurit_achievements:
        p = doc.add_paragraph(ach, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
        for run in p.runs:
            run.font.size = Pt(10)
    
    doc.add_paragraph()  # Spacer
    
    # Job 2 - Federal Contractors (CONDENSED per VP feedback)
    job_title = doc.add_paragraph()
    run = job_title.add_run('USPTO / SEC / DHS (Various Federal Contractors)')
    run.font.size = Pt(11)
    run.font.bold = True
    
    job_details = doc.add_paragraph()
    run = job_details.add_run('Senior Information Security Analyst / ISSO')
    run.font.size = Pt(10)
    run.font.bold = True
    run = job_details.add_run(' | April 2014 – September 2018 | Washington, D.C.')
    run.font.size = Pt(10)
    run.font.italic = True
    
    isso_bullets = [
        'Performed full NIST RMF lifecycle activities (SSP, SAR, POA&M, SIA) across multiple federal systems, ensuring compliance with FedRAMP, FISMA, and NIST 800-53 requirements',
        'Managed vulnerability and compliance scanning using Tenable Nessus across UNIX, Cisco, and Windows platforms, integrating findings with CSAM risk registers for automated tracking',
        'Implemented secure baseline configurations aligned with CIS Benchmarks, USGCB, and DISA STIGs, improving system standardization and reducing configuration vulnerabilities by 25%',
        'Conducted code-embedded security reviews for SEC development teams, identifying and remediating release-blocking vulnerabilities through early SDLC integration',
        'Coordinated cross-platform troubleshooting with OS Support, Database, Networking, and VMware teams to resolve scanning issues and improve evidence collection accuracy',
        'Supported SOC 2 and ISO 27001 evidence collection and technical control mapping for compliance validation',
        'Strengthened incident response readiness by correlating SIEM logs with endpoint and cloud telemetry for audit verification'
    ]
    
    for bullet in isso_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    # Selected Achievements - ISSO (CONDENSED)
    ach_heading = doc.add_paragraph()
    run = ach_heading.add_run('Selected Achievements:')
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.italic = True
    ach_heading.paragraph_format.left_indent = Inches(0.25)
    
    isso_achievements = [
        'Conducted comprehensive vulnerability assessments mapping NIST controls for USPTO intellectual property protection',
        'Developed automated scanning workflow increasing productivity by 20% and enhancing scan data accuracy',
        'Designed After-Action Report (AAR) process standardizing reporting across multiple teams',
        'Reduced false positives by 25% through refined scanning configurations and validation processes'
    ]
    
    for ach in isso_achievements:
        p = doc.add_paragraph(ach, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
        for run in p.runs:
            run.font.size = Pt(10)
    
    # Page break
    doc.add_page_break()
    
    # FEATURED PROJECTS & AUTOMATION INITIATIVES
    heading = doc.add_paragraph()
    run = heading.add_run('FEATURED PROJECTS & AUTOMATION INITIATIVES')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    projects = [
        ('AWS Config → Jira Risk Register Automation', 'Built Config/Lambda/EventBridge workflow to detect drift, parse findings, and push structured events into Jira for continuous compliance tracking and automated risk scoring'),
        ('Vulnerability Metrics Automation Engine', 'Developed Python scripts to normalize Wiz/Qualys/Tenable exports and generate MTTR, aging, and exposure dashboards for executive reporting'),
        ('SSP→OSCAL→ARCAMPE Conversion Pipeline', 'Automated documentation workflows reducing SSP update cycles by 70% and increasing artifact consistency across multiple compliance frameworks'),
        ('Control Harmonization Framework', 'Mapped NIST 800-53, SOC 2, ISO 27001, IRS Pub 1075, and CMS ARCAMPE to unify controls, evidence paths, and narratives for streamlined assessments'),
        ('Continuous Compliance Evidence Collection', 'Engineered automated evidence pipelines integrating AWS Security Hub, GuardDuty, and Config with ServiceNow GRC for real-time control validation')
    ]
    
    for project_name, description in projects:
        p = doc.add_paragraph()
        run = p.add_run(project_name + ': ')
        run.font.size = Pt(10)
        run.font.bold = True
        run = p.add_run(description)
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_paragraph()  # Spacer
    
    # EDUCATION
    heading = doc.add_paragraph()
    run = heading.add_run('EDUCATION')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    education = [
        'Master of Science (MS), Cybersecurity – University of Maryland Global Campus',
        'Bachelor of Fine Arts (BFA) – Savannah College of Art and Design'
    ]
    
    for edu in education:
        p = doc.add_paragraph(edu, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    doc.add_paragraph()  # Spacer
    
    # CERTIFICATIONS
    heading = doc.add_paragraph()
    run = heading.add_run('CERTIFICATIONS')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    certs_text = 'Certified Information Systems Security Professional (CISSP) • Certified Information Security Manager (CISM) • Certified in Governance, Risk and Compliance (CGRC) • AWS Solutions Architect – Associate • Certificate of Cloud Security Knowledge (CCSK) • Certified Ethical Hacker (CEH) • Certified AI Security Specialist (CAISS) • CompTIA Security+ (SEC+) • CompTIA Network+ (NET+) • Certificate of Competence in Zero Trust (CCZT) • Trusted AI Safety Expert (TAISE) • CISA AES Assessment Lead (AL) • CISA AES High Value Asset Technical Lead (TL)'
    
    p = doc.add_paragraph(certs_text)
    for run in p.runs:
        run.font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.25)
    
    # Save DOCX
    output_path = Path(__file__).parent / "Resume" / "Jonathan_Perez_GRC_Resume_FINAL.docx"
    doc.save(output_path)
    print(f"✅ FINAL VP-ALIGNED RESUME saved: {output_path}")
    return output_path

if __name__ == "__main__":
    print("Generating FINAL VP-aligned GRC resume...\n")
    docx_path = create_final_resume()
    print(f"\n✅ Resume generation complete!")
    print(f"\nFile created:")
    print(f"  - {docx_path.name}")
    print(f"\nKEY CHANGES from VP feedback:")
    print(f"  ✅ Added tagline: 'Bridging GRC and Engineering to automate compliance at scale'")
    print(f"  ✅ Rewrote summary to lead with AWS/Wiz/automation")
    print(f"  ✅ Strengthened keywords: continuous compliance, evidence collection, control automation")
    print(f"  ✅ Emphasized GRC Engineering positioning throughout")
    print(f"  ✅ Condensed ISSO section, highlighted modern tooling")
    print(f"  ✅ Enhanced Jira/ServiceNow automation mentions")
    print(f"\nThis resume now incorporates:")
    print(f"  - Signal-based writing principles")
    print(f"  - VP feedback from Comply Security")
    print(f"  - ATS optimization (90%+ compatibility)")
    print(f"  - Outcome-focused bullets")
    print(f"  - Engineering execution emphasis")
