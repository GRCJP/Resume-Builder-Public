#!/usr/bin/env python3
"""Generate clean, properly formatted GRC resume in DOCX"""

import sys
from pathlib import Path

def install_packages():
    """Install required packages"""
    import subprocess
    try:
        __import__('docx')
    except ImportError:
        print("Installing python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

install_packages()

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def add_heading(doc, text, level=1):
    """Add a formatted heading"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12 if level == 1 else 11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.space_before = Pt(6 if level == 1 else 3)
    p.space_after = Pt(3)
    return p

def add_bullet(doc, text, indent=0.25):
    """Add a bullet point"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def add_text(doc, text, bold=False, size=10):
    """Add regular text"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    p.paragraph_format.space_after = Pt(3)
    return p

def create_clean_resume():
    """Create clean, professional DOCX resume"""
    doc = Document()
    
    # Set narrow margins for 2-page fit
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    # HEADER
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run('JONATHAN L. PEREZ, MS, CISSP, CISM')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    name.space_after = Pt(2)
    
    # Tagline
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run('Bridging GRC and Engineering to automate compliance at scale')
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    tagline.space_after = Pt(2)
    
    # Contact
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact.add_run('Aldie, VA | 703-843-6887 | jleepe@outlook.com | linkedin.com/in/cyberjp | securitybyjp.com')
    run.font.size = Pt(9)
    contact.space_after = Pt(2)
    
    # Clearance
    clearance = doc.add_paragraph()
    clearance.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = clearance.add_run('Active Secret Clearance')
    run.font.size = Pt(9)
    run.font.italic = True
    clearance.space_after = Pt(6)
    
    # PROFESSIONAL SUMMARY
    add_heading(doc, 'PROFESSIONAL SUMMARY')
    summary = doc.add_paragraph()
    run = summary.add_run('GRC Engineer leveraging AWS Config, Lambda, EventBridge, and Security Hub to automate continuous compliance and evidence collection across hybrid cloud environments. 10+ years bridging NIST RMF, FedRAMP, SOC 2, ISO 27001, and PCI DSS governance with engineering execution using Wiz, Qualys, Tenable, and Python automation. Proven track record reducing vulnerabilities 83% and cutting documentation cycles 70% through control automation and risk register integration.')
    run.font.size = Pt(10)
    summary.paragraph_format.space_after = Pt(6)
    
    # CORE COMPETENCIES
    add_heading(doc, 'CORE COMPETENCIES')
    
    competencies = [
        ('Cloud & Automation:', 'AWS Config, Lambda, EventBridge, Security Hub, Systems Manager, GuardDuty, Python, OSCAL, CI/CD'),
        ('GRC Platforms:', 'ServiceNow GRC, DOJ CSAM, RegScale, Jira Automation, API-based evidence pipelines'),
        ('Vulnerability Tools:', 'Wiz, Qualys, Tenable, CrowdStrike Falcon, CVSSv4, EPSS, MITRE ATT&CK'),
        ('Frameworks:', 'NIST RMF, 800-53/171/30, FedRAMP, FISMA, SOC 2, ISO 27001, PCI DSS 4.0, HIPAA, GDPR'),
        ('Security Operations:', 'SIEM (Splunk, Security Hub, GuardDuty), incident response, continuous monitoring')
    ]
    
    for category, details in competencies:
        p = doc.add_paragraph()
        run = p.add_run(category + ' ')
        run.font.size = Pt(9)
        run.font.bold = True
        run = p.add_run(details)
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
    
    doc.add_paragraph().space_after = Pt(3)
    
    # KEY ACHIEVEMENTS
    add_heading(doc, 'KEY ACHIEVEMENTS')
    achievements = [
        'Reduced enterprise vulnerabilities from 70,000 to 12,000 across 10,000+ assets (83% reduction)',
        'Engineered AWS-based continuous compliance pipelines reducing manual documentation cycles by 70%',
        'Automated vulnerability-to-POA&M workflows cutting manual review hours by 30%',
        'Achieved 90%+ CIS/STIG compliance across 7,500+ cloud assets',
        'Reduced critical vulnerabilities by 65% in one quarter through automated remediation'
    ]
    for ach in achievements:
        add_bullet(doc, ach)
    
    doc.add_paragraph().space_after = Pt(3)
    
    # PROFESSIONAL EXPERIENCE
    add_heading(doc, 'PROFESSIONAL EXPERIENCE')
    
    # Job 1
    job = doc.add_paragraph()
    run = job.add_run('ASSURIT')
    run.font.size = Pt(10)
    run.font.bold = True
    job.space_after = Pt(1)
    
    title = doc.add_paragraph()
    run = title.add_run('Senior Cybersecurity Specialist / Deputy Project Manager | ')
    run.font.size = Pt(10)
    run.font.bold = True
    run = title.add_run('September 2018 - Present | Washington, D.C.')
    run.font.size = Pt(10)
    run.font.italic = True
    title.space_after = Pt(2)
    
    bullets = [
        'Lead enterprise vulnerability management program across 10,000+ assets, integrating Wiz, Qualys, and Tenable with Jira and ServiceNow for automated remediation tracking',
        'Engineer continuous compliance pipelines using AWS Config, Lambda, and EventBridge to automatically collect evidence and detect configuration drift in real-time',
        'Automate RMF evidence collection transforming scanner data into normalized, control-mapped findings, reducing manual review time by 30%',
        'Serve as GRC Engineering SME on NIST RMF, FedRAMP, SOC 2, ISO 27001, and PCI DSS compliance, bridging governance requirements with automated technical implementation',
        'Develop control automation workflows using Python and RegScale, converting SSP to OSCAL to ARCAMPE documentation and reducing compliance cycles by 70%',
        'Lead security assessments for MD THINK, MD Benefits, FEC, and USPTO across NIST 800-53/171, FedRAMP, ISO 27001, SOC 2, and IRS Pub 1075',
        'Integrate SIEM data from Security Hub, GuardDuty, and Splunk with GRC platforms for automated evidence collection and control validation',
        'Build Jira-based risk register automation improving risk scoring, tracking workflows, and continuous monitoring updates',
        'Validate CI/CD pipeline security controls including ECR scanning, Wiz SCA, and SSM patching aligned with NIST 800-53 and FedRAMP baselines',
        'Deliver ISSO enablement programs reducing artifact submission errors by 40% through standardized workflows and training',
        'Drive weekly vulnerability management sessions prioritizing findings by CVSS/EPSS and asset criticality, coordinating automated remediation',
        'Manage POA&M lifecycle from identification through closure using automated tracking for timely audit finding resolution'
    ]
    for bullet in bullets:
        add_bullet(doc, bullet)
    
    doc.add_paragraph().space_after = Pt(3)
    
    # Job 2
    job = doc.add_paragraph()
    run = job.add_run('USPTO / SEC / DHS (Various Federal Contractors)')
    run.font.size = Pt(10)
    run.font.bold = True
    job.space_after = Pt(1)
    
    title = doc.add_paragraph()
    run = title.add_run('Senior Information Security Analyst / ISSO | ')
    run.font.size = Pt(10)
    run.font.bold = True
    run = title.add_run('April 2014 - September 2018 | Washington, D.C.')
    run.font.size = Pt(10)
    run.font.italic = True
    title.space_after = Pt(2)
    
    bullets = [
        'Performed full NIST RMF lifecycle activities (SSP, SAR, POA&M, SIA) ensuring compliance with FedRAMP, FISMA, and NIST 800-53',
        'Managed vulnerability scanning using Tenable Nessus across UNIX, Cisco, and Windows platforms, integrating findings with CSAM risk registers',
        'Implemented secure baseline configurations aligned with CIS Benchmarks, USGCB, and DISA STIGs, reducing configuration vulnerabilities by 25%',
        'Conducted code-embedded security reviews for SEC development teams, identifying and remediating release-blocking vulnerabilities',
        'Supported SOC 2 and ISO 27001 evidence collection and technical control mapping for compliance validation',
        'Developed automated scanning workflow increasing productivity by 20% and enhancing scan data accuracy',
        'Designed After-Action Report process standardizing reporting across multiple teams'
    ]
    for bullet in bullets:
        add_bullet(doc, bullet)
    
    # Page break
    doc.add_page_break()
    
    # FEATURED PROJECTS
    add_heading(doc, 'FEATURED PROJECTS & AUTOMATION INITIATIVES')
    
    projects = [
        ('AWS Config to Jira Risk Register Automation:', 'Built Config/Lambda/EventBridge workflow detecting drift and pushing structured events to Jira for continuous compliance tracking'),
        ('Vulnerability Metrics Automation Engine:', 'Developed Python scripts normalizing Wiz/Qualys/Tenable exports and generating MTTR, aging, and exposure dashboards'),
        ('SSP to OSCAL to ARCAMPE Conversion Pipeline:', 'Automated documentation workflows reducing update cycles by 70% and increasing artifact consistency'),
        ('Control Harmonization Framework:', 'Mapped NIST 800-53, SOC 2, ISO 27001, IRS Pub 1075, and CMS ARCAMPE to unify controls and evidence paths'),
        ('Continuous Compliance Evidence Collection:', 'Engineered automated pipelines integrating AWS Security Hub and GuardDuty with ServiceNow GRC for real-time validation')
    ]
    
    for project_name, description in projects:
        p = doc.add_paragraph()
        run = p.add_run(project_name + ' ')
        run.font.size = Pt(9)
        run.font.bold = True
        run = p.add_run(description)
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
    
    doc.add_paragraph().space_after = Pt(3)
    
    # EDUCATION
    add_heading(doc, 'EDUCATION')
    add_bullet(doc, 'Master of Science (MS), Cybersecurity - University of Maryland Global Campus')
    add_bullet(doc, 'Bachelor of Fine Arts (BFA) - Savannah College of Art and Design')
    
    doc.add_paragraph().space_after = Pt(3)
    
    # CERTIFICATIONS
    add_heading(doc, 'CERTIFICATIONS')
    p = doc.add_paragraph()
    certs = 'CISSP | CISM | CGRC | AWS Solutions Architect - Associate | CCSK | CEH | CAISS | Security+ | Network+ | CCZT | TAISE | CISA AES Assessment Lead (AL) | CISA AES Technical Lead (TL)'
    run = p.add_run(certs)
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    
    # Save
    output_path = Path(__file__).parent / "Resume" / "Jonathan_Perez_GRC_Resume_CLEAN.docx"
    doc.save(output_path)
    print(f"✅ Clean resume saved: {output_path}")
    print(f"\nFormatting improvements:")
    print(f"  ✅ Professional spacing (2 pages)")
    print(f"  ✅ Clean bullets (no arrows/slashes)")
    print(f"  ✅ Highlighted tools at top")
    print(f"  ✅ Readable layout")
    print(f"  ✅ Proper indentation")
    return output_path

if __name__ == "__main__":
    print("Generating clean, professional GRC resume...\n")
    create_clean_resume()
    print(f"\n✅ Done! Open the file in Microsoft Word.")
