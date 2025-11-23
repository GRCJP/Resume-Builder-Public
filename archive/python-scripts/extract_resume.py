#!/usr/bin/env python3
"""Extract text from resume files for analysis"""

import sys
from pathlib import Path

def extract_from_docx(file_path):
    """Extract text from .docx file"""
    try:
        from docx import Document
        doc = Document(file_path)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        return '\n'.join(text)
    except ImportError:
        print("Installing python-docx...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        from docx import Document
        doc = Document(file_path)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        return '\n'.join(text)

def extract_from_pdf(file_path):
    """Extract text from .pdf file"""
    try:
        import PyPDF2
        text = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text.append(page.extract_text())
        return '\n'.join(text)
    except ImportError:
        print("Installing PyPDF2...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "PyPDF2"])
        import PyPDF2
        text = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text.append(page.extract_text())
        return '\n'.join(text)

if __name__ == "__main__":
    resume_dir = Path(__file__).parent / "Resume"
    
    # Extract from DOCX
    docx_file = resume_dir / "Jonathan _GRC Resume Update.docx"
    if docx_file.exists():
        print("=== UPDATED RESUME (DOCX) ===\n")
        docx_text = extract_from_docx(docx_file)
        print(docx_text)
        
        # Save to text file
        with open(resume_dir / "updated_resume.txt", "w") as f:
            f.write(docx_text)
    
    print("\n\n" + "="*80 + "\n\n")
    
    # Extract from PDF
    pdf_file = resume_dir / "Jonathan_Perez_GRC_2025 .pdf"
    if pdf_file.exists():
        print("=== ORIGINAL RESUME (PDF) ===\n")
        pdf_text = extract_from_pdf(pdf_file)
        print(pdf_text)
        
        # Save to text file
        with open(resume_dir / "original_resume.txt", "w") as f:
            f.write(pdf_text)
    
    print("\n\nText files saved to Resume/ folder")
