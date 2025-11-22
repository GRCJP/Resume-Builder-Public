#!/usr/bin/env python3
"""Generate VP-aligned GRC resume matching template format"""

import sys
from pathlib import Path

def install_packages():
    """Install required packages"""
    import subprocess
    try:
        __import__('docx')
    except ImportError:
        print("Installing python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

install_packages()

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_section_heading(doc, text):
    """Add a section heading matching template style"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_bullet(doc, text):
    """Add a bullet point"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def add_text(doc, text, bold=False, size=10, italic=False):
    """Add regular text"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    p.paragraph_format.space_after = Pt(3)
    return p

def create_resume():
    """Create VP-aligned resume"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    # HEADER - Name
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run('JONATHAN L. PEREZ, MS, CISSP, CISM')
    run.font.size = Pt(14)
    run.font.bold = True
    name.space_after = Pt(3)
    
    # Contact Info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact.add_run('Aldie, VA | 703-843-6887 | jleepe@outlook.com | linkedin.com/in/cyberjp | securitybyjp.com')
    run.font.size = Pt(9)
    contact.space_after = Pt(8)
    
    # PROFESSIONAL SUMMARY (Heading 1 style from template)
    summary_heading = doc.add_paragraph()
    summary_heading.paragraph_format.space_before = Pt(6)
    summary_heading.paragraph_format.space_after = Pt(4)
    run = summary_heading.add_run('GRC ENGINEER & SECURITY COMPLIANCE SME')
    run.font.size = Pt(11)
    run.font.bold = True
    
    # Summary content
    summary = doc.add_paragraph()
    run = summary.add_run('GRC Engineer leveraging AWS Config, Lambda, EventBridge, and Security Hub to automate continuous compliance and evidence collection across hybrid cloud environments. 10+ years bridging NIST RMF, FedRAMP, SOC 2, ISO 27001, and PCI DSS governance with engineering execution using Wiz, Qualys, Tenable, and Python automation. Proven track record reducing vulnerabilities 83% and cutting documentation cycles 70% through control automation and risk register integration.')
    run.font.size = Pt(10)
    summary.paragraph_format.space_after = Pt(8)
    
    # CERTIFICATIONS & CLEARANCE
    add_section_heading(doc, 'CERTIFICATIONS & SECURITY CLEARANCE')
    
    certs_p1 = doc.add_paragraph()
    run = certs_p1.add_run('ISC2 - CISSP | ISC2 - CISM | ISACA - CGRC | AWS Solutions Architect - Associate | CSA - CCSK | EC-Council - CEH')
    run.font.size = Pt(9)
    certs_p1.paragraph_format.space_after = Pt(2)
    
    certs_p2 = doc.add_paragraph()
    run = certs_p2.add_run('CompTIA Security+ | CompTIA Network+ | CSA - CCZT | CAISS | TAISE | CISA AES Assessment Lead (AL) | CISA AES Technical Lead (TL)')
    run.font.size = Pt(9)
    certs_p2.paragraph_format.space_after = Pt(2)
    
    clearance = doc.add_paragraph()
    run = clearance.add_run('Active Secret Clearance')
    run.font.size = Pt(9)
    run.font.bold = True
    clearance.paragraph_format.space_after = Pt(8)
    
    # EDUCATION
    add_section_heading(doc, 'EDUCATION')
    
    edu1 = doc.add_paragraph()
    run = edu1.add_run('University of Maryland Global Campus – ')
    run.font.size = Pt(10)
    run = edu1.add_run('Master of Science (MS), Cybersecurity')
    run.font.size = Pt(10)
    edu1.paragraph_format.space_after = Pt(2)
    
    edu2 = doc.add_paragraph()
    run = edu2.add_run('Savannah College of Art and Design – ')
    run.font.size = Pt(10)
    run = edu2.add_run('Bachelor of Fine Arts (BFA)')
    run.font.size = Pt(10)
    edu2.paragraph_format.space_after = Pt(8)
    
    # TOOLS & PLATFORMS (VP recommended instead of Core Competencies)
    add_section_heading(doc, 'TOOLS & PLATFORMS')
    
    tools = [
        ('Cloud & Automation:', 'AWS Config, Lambda, EventBridge, Security Hub, Systems Manager, GuardDuty, Python, OSCAL, CI/CD'),
        ('GRC Platforms:', 'ServiceNow GRC, DOJ CSAM, RegScale, Jira Automation, API-based evidence pipelines'),
        ('Vulnerability Management:', 'Wiz, Qualys VMDR, Tenable Nessus, CrowdStrike Falcon, CVSSv4, EPSS, MITRE ATT&CK'),
        ('Compliance Frameworks:', 'NIST RMF, 800-53/171/30, FedRAMP, FISMA, SOC 2, ISO 27001, PCI DSS 4.0, HIPAA, GDPR, IRS Pub 1075'),
        ('Security Operations:', 'SIEM (Splunk, AWS Security Hub, GuardDuty), incident response, continuous monitoring, threat intelligence')
    ]
    
    for category, details in tools:
        p = doc.add_paragraph()
        run = p.add_run(category + ' ')
        run.font.size = Pt(9)
        run.font.bold = True
        run = p.add_run(details)
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
    
    doc.add_paragraph().space_after = Pt(6)
    
    # PROFESSIONAL EXPERIENCE
    add_section_heading(doc, 'PROFESSIONAL EXPERIENCE')
    
    # Job 1 - ASSURIT (with updated title)
    job1_header = doc.add_paragraph()
    run = job1_header.add_run('Assurit – Washington, D.C. (Supporting Federal Agencies) – ')
    run.font.size = Pt(10)
    run.font.bold = True
    run = job1_header.add_run('GRC Engineer / Senior Cybersecurity Specialist')
    run.font.size = Pt(10)
    run.font.bold = True
    job1_header.paragraph_format.space_after = Pt(1)
    
    job1_date = doc.add_paragraph()
    run = job1_date.add_run('September 2018 - Present')
    run.font.size = Pt(9)
    run.font.italic = True
    job1_date.paragraph_format.space_after = Pt(3)
    
    # Bullets emphasizing engineering and automation
    bullets1 = [
        'Lead enterprise vulnerability management program across 10,000+ assets, reducing vulnerabilities from 70,000 to 12,000 (83% reduction) through automated remediation workflows integrating Wiz, Qualys, and Tenable with Jira and ServiceNow',
        'Engineer continuous compliance pipelines using AWS Config, Lambda, and EventBridge to automatically collect evidence and detect configuration drift in real-time, reducing manual documentation cycles by 70%',
        'Automate control testing and evidence collection transforming scanner data into normalized, control-mapped findings, reducing manual review time by 30% and improving audit readiness',
        'Serve as GRC Engineering SME bridging NIST RMF, FedRAMP, SOC 2, ISO 27001, and PCI DSS governance requirements with automated technical implementation using Python and RegScale',
        'Develop control automation workflows converting SSP to OSCAL to ARCAMPE documentation formats, enabling continuous compliance monitoring and reducing update cycles by 70%',
        'Lead security assessments for MD THINK, MD Benefits, FEC, and USPTO across NIST 800-53/171, FedRAMP, ISO 27001, SOC 2, and IRS Pub 1075 frameworks',
        'Integrate SIEM data from AWS Security Hub, GuardDuty, and Splunk with GRC platforms for automated evidence collection and real-time control validation',
        'Build Jira-based risk register automation improving risk scoring accuracy, tracking workflows, and continuous monitoring updates across multiple compliance programs',
        'Validate CI/CD pipeline security controls including ECR scanning, Wiz SCA, and SSM patching aligned with NIST 800-53 and FedRAMP baselines',
        'Drive weekly vulnerability management sessions prioritizing findings by CVSS/EPSS and asset criticality, coordinating automated remediation across cloud and on-premise environments',
        'Manage POA&M lifecycle from identification through closure using automated tracking, ensuring timely resolution of audit findings and maintaining compliance posture',
        'Deliver ISSO enablement programs reducing artifact submission errors by 40% through standardized workflows, templates, and training'
    ]
    
    for bullet in bullets1:
        add_bullet(doc, bullet)
    
    doc.add_paragraph().space_after = Pt(6)
    
    # Job 2 - Federal Roles (condensed per VP feedback)
    job2_header = doc.add_paragraph()
    run = job2_header.add_run('USPTO / SEC / DHS (Various Federal Contractors) – ')
    run.font.size = Pt(10)
    run.font.bold = True
    run = job2_header.add_run('Senior Information Security Analyst / ISSO')
    run.font.size = Pt(10)
    run.font.bold = True
    job2_header.paragraph_format.space_after = Pt(1)
    
    job2_date = doc.add_paragraph()
    run = job2_date.add_run('April 2014 - September 2018')
    run.font.size = Pt(9)
    run.font.italic = True
    job2_date.paragraph_format.space_after = Pt(3)
    
    # Condensed bullets focusing on outcomes
    bullets2 = [
        'Performed full NIST RMF lifecycle activities (SSP, SAR, POA&M, SIA) ensuring compliance with FedRAMP, FISMA, and NIST 800-53 across multiple federal systems',
        'Managed vulnerability scanning using Tenable Nessus across UNIX, Cisco, and Windows platforms, integrating findings with CSAM risk registers and reducing vulnerabilities by 25%',
        'Implemented secure baseline configurations aligned with CIS Benchmarks, USGCB, and DISA STIGs, improving security posture and reducing configuration-based vulnerabilities',
        'Conducted code-embedded security reviews for SEC development teams, identifying and remediating release-blocking vulnerabilities before production deployment',
        'Supported SOC 2 and ISO 27001 evidence collection and technical control mapping for compliance validation and audit readiness',
        'Developed automated scanning workflows increasing productivity by 20% and enhancing scan data accuracy across enterprise environments'
    ]
    
    for bullet in bullets2:
        add_bullet(doc, bullet)
    
    # Save
    output_path = Path(__file__).parent / "Resume" / "Updated Resumes" / "Jonathan_Perez_GRC_Master.docx"
    doc.save(output_path)
    print(f"✅ VP-aligned resume saved: {output_path}")
    print(f"\nKey improvements:")
    print(f"  ✅ Section headings match template style")
    print(f"  ✅ 'TOOLS & PLATFORMS' instead of 'Core Competencies'")
    print(f"  ✅ Job title: 'GRC Engineer / Senior Cybersecurity Specialist'")
    print(f"  ✅ Summary emphasizes AWS/Wiz/automation (VP feedback)")
    print(f"  ✅ Condensed federal roles per VP recommendation")
    print(f"  ✅ Engineering-focused language throughout")
    return output_path

if __name__ == "__main__":
    print("Generating VP-aligned GRC resume with template formatting...\n")
    create_resume()
    print(f"\n✅ Done! Open the file in Microsoft Word to review.")
